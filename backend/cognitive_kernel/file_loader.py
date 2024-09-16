from typing import List
import docx
import fitz
import json
import re


def split_text_into_chunks(text, max_chars=300):
    chunks = []
    current_chunk = ""

    for char in text:
        if len(current_chunk) + 1 <= max_chars:
            current_chunk += char
        else:
            if char == "。":
                current_chunk += char
                chunks.append(current_chunk)
                current_chunk = ""
            else:
                # Find the last complete sentence to split
                last_period_idx = current_chunk.rfind("。")
                if last_period_idx != -1:
                    # Split at the last complete sentence
                    chunks.append(current_chunk[: last_period_idx + 1])
                    current_chunk = current_chunk[last_period_idx + 1 :] + char
                else:
                    chunks.append(current_chunk)
                    current_chunk = char

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def preprocessing_text(text: str) -> str:

    text = re.sub(r" +", " ", text)
    text = text.strip().strip("\n")

    return text


def paragraph_to_chunk(paragraph: str, chunk_num: int = 5) -> List[str]:

    paragraph = preprocessing_text(paragraph)
    sentences = split_text_into_chunks(paragraph)

    batches = [
        sentences[i : i + chunk_num] for i in range(0, len(sentences), chunk_num)
    ]

    chunks = list()
    for batch in batches:
        chunks.append(" ".join(batch))

    return chunks


def read_file_content(
    file_location: str, file_name: str, chunk_num: int = 5
) -> List[str]:
    """Reads an uploaded file from a given location and returns its content as a list of strings."""

    if file_name.endswith(".docx"):
        return read_docx(file_location, chunk_num)
    elif file_name.endswith(".pdf"):
        return read_pdf(file_location, chunk_num)
    elif file_name.endswith(".txt"):
        return read_txt(file_location, chunk_num)
    else:
        raise NotImplementedError(f"File type {file_name.split('.')[-1]} not supported")


def read_docx(path: str, chunk_num: int = 5) -> List[str]:
    doc = docx.Document(path)
    contents = []
    for paragraph in doc.paragraphs:
        if paragraph.text:

            chunks = paragraph_to_chunk(paragraph.text, chunk_num)
            for chunk in chunks:
                contents.append(
                    {
                        "text": chunk,
                        "meta": json.dumps(
                            {"Row Number": len(contents) + 1}, ensure_ascii=False
                        ),
                    }
                )

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join([cell.text for cell in row.cells if cell.text])
            contents.append(
                {
                    "text": row_text,
                    "meta": json.dumps(
                        {"Row Number": len(contents) + 1}, ensure_ascii=False
                    ),
                }
            )
    return contents


def read_pdf(path: str, chunk_num: int = 5) -> List[str]:
    contents = []
    with fitz.open(path) as doc:
        title = doc.metadata.get("title")
        for idx, page in enumerate(doc):
            imagepath = path + f".page_{idx+1}.png"

            chunks = paragraph_to_chunk(page.get_text(), chunk_num)
            for chunk in chunks:
                contents.append(
                    {
                        "text": chunk,
                        "meta": json.dumps(
                            {"Title": title, "Page Number": idx + 1},
                            ensure_ascii=False,
                        ),
                    }
                )
    return contents


def read_txt(path: str, chunk_num: int = 5) -> List[str]:
    contents = []
    with open(path, "r", encoding="utf-8") as file:
        idx = 0
        for line in file:
            chunks = paragraph_to_chunk(line, chunk_num)
            for chunk in chunks:
                contents.append(
                    {
                        "text": chunk,
                        "meta": json.dumps(
                            {"Paragraph Number": idx + 1}, ensure_ascii=False
                        ),
                    }
                )
            idx += 1
    return contents
